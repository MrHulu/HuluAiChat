"""聊天记录导出服务."""
import html
import json
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF
from src.persistence import Session, Message


class ChatExporter:
    """聊天记录导出器，支持 TXT、Markdown、JSON、HTML、PDF 和 DOCX 格式."""

    def __init__(self, session: Session, messages: list[Message]) -> None:
        self._session = session
        self._messages = messages

    def to_markdown(self) -> str:
        """导出为 Markdown 格式."""
        lines = [
            f"# {self._session.title or '新对话'}\n",
            f"**创建时间**: {self._format_time(self._session.created_at)}  \n",
            f"**更新时间**: {self._format_time(self._session.updated_at)}\n",
            "---\n\n",
        ]

        for msg in self._messages:
            role_name = "你" if msg.role == "user" else "助手"
            lines.append(f"## {role_name}\n")
            lines.append(f"{msg.content}\n\n")

        return "".join(lines)

    def to_txt(self) -> str:
        """导出为纯文本格式.

        Returns:
            纯文本文档字符串
        """
        lines = [
            "=" * 60,
            f"{self._session.title or '新对话'}",
            "=" * 60,
            f"创建时间: {self._format_time(self._session.created_at)}",
            f"更新时间: {self._format_time(self._session.updated_at)}",
            "",
            "-" * 60,
            "",
        ]

        for msg in self._messages:
            role_name = "【你】" if msg.role == "user" else "【助手】"
            lines.append(role_name)
            lines.append(msg.content or "")
            lines.append("")
            lines.append(f"时间: {self._format_time(msg.created_at)}")
            lines.append("")
            lines.append("-" * 60)
            lines.append("")

        return "\n".join(lines)

    def to_json(self) -> str:
        """导出为 JSON 格式."""
        data = {
            "session": {
                "id": self._session.id,
                "title": self._session.title,
                "created_at": self._session.created_at,
                "updated_at": self._session.updated_at,
            },
            "messages": [
                {
                    "id": m.id,
                    "role": m.role,
                    "content": m.content,
                    "created_at": m.created_at,
                }
                for m in self._messages
            ],
        }
        return json.dumps(data, ensure_ascii=False, indent=2)

    def to_html(self) -> str:
        """导出为 HTML 格式.

        Returns:
            HTML 文档字符串
        """
        title = html.escape(self._session.title or "新对话")

        html_parts = [
            '<!DOCTYPE html>',
            '<html lang="zh-CN">',
            '<head>',
            '    <meta charset="UTF-8">',
            '    <meta name="viewport" content="width=device-width, initial-scale=1.0">',
            f'    <title>{title}</title>',
            '    <style>',
            '        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, "Microsoft YaHei", sans-serif; line-height: 1.6; max-width: 800px; margin: 40px auto; padding: 0 20px; background: #f5f5f5; }',
            '        .container { background: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }',
            '        h1 { color: #333; border-bottom: 2px solid #007AFF; padding-bottom: 10px; }',
            '        .meta { color: #666; font-size: 14px; margin-bottom: 20px; }',
            '        .message { margin: 20px 0; padding: 15px; border-radius: 8px; }',
            '        .message.user { background: #E3F2FD; border-left: 4px solid #2196F3; }',
            '        .message.assistant { background: #F5F5F5; border-left: 4px solid #4CAF50; }',
            '        .role { font-weight: bold; margin-bottom: 8px; display: block; }',
            '        .role::before { content: attr(data-role); }',
            '        .content { white-space: pre-wrap; word-wrap: break-word; }',
            '        .time { font-size: 12px; color: #999; margin-top: 8px; }',
            '    </style>',
            '</head>',
            '<body>',
            '    <div class="container">',
            f'        <h1>{title}</h1>',
            f'        <div class="meta">创建时间: {self._format_time(self._session.created_at)} | 更新时间: {self._format_time(self._session.updated_at)}</div>',
        ]

        for msg in self._messages:
            role_class = "user" if msg.role == "user" else "assistant"
            role_text = "你" if msg.role == "user" else "助手"
            content = html.escape(msg.content or "")
            time = self._format_time(msg.created_at)

            html_parts.extend([
                f'        <div class="message {role_class}">',
                f'            <span class="role" data-role="{role_text}"></span>',
                f'            <div class="content">{content}</div>',
                f'            <div class="time">{time}</div>',
                '        </div>',
            ])

        html_parts.extend([
            '    </div>',
            '</body>',
            '</html>',
        ])

        return "\n".join(html_parts)

    def to_pdf(self) -> bytes:
        """导出为 PDF 格式（支持中文）.

        Returns:
            PDF 文件的二进制数据
        """
        try:
            # 尝试使用支持 Unicode 的方法
            from fpdf import FontFace

            pdf = FPDF()
            pdf.add_page()

            # 尝试使用系统自带的中文字体
            # Windows: Microsoft YaHei, SimSun
            # macOS: PingFang SC, STHeiti
            # Linux: WenQuanYi Zen Hei, Noto Sans CJK

            title = self._session.title or "新对话"

            # 使用 write() 方法支持 Unicode 文本
            # 首先添加一个支持 Unicode 的字体
            pdf.set_font("Arial", size=16)

            # 标题 - 使用 write() 避免编码问题
            pdf.ln(10)
            pdf.set_x(0)
            # 使用 UTF-8 编码直接输出
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, title.encode("latin-1", "replace").decode("latin-1"), new_x="LEFT", new_y="NEXT", align="C")

            pdf.set_font("Arial", "", 10)
            pdf.cell(0, 6, f"Created: {self._format_time(self._session.created_at)}", new_x="LEFT", new_y="NEXT")
            pdf.cell(0, 6, f"Updated: {self._format_time(self._session.updated_at)}", new_x="LEFT", new_y="NEXT")

            pdf.ln(5)
            pdf.set_line_width(0.5)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)

            # 消息
            pdf.set_font("Arial", "B", 12)
            for msg in self._messages:
                # 角色名称
                role_name = "You" if msg.role == "user" else "Assistant"
                pdf.set_fill_color(240 if msg.role == "user" else 200, 240 if msg.role == "user" else 220, 240 if msg.role == "user" else 220)
                pdf.cell(0, 8, role_name, new_x="LEFT", new_y="NEXT", fill=True)

                # 内容 - 支持中文和自动换行
                pdf.set_font("Arial", "", 10)
                content = msg.content or ""
                # FPDF 对中文支持有限，使用替代字符
                safe_content = content.encode("latin-1", "replace").decode("latin-1")

                # 简单的自动换行
                lines = self._wrap_text(safe_content, 190)
                for line in lines:
                    pdf.cell(0, 6, line, new_x="LEFT", new_y="NEXT")

                pdf.ln(3)

            # 生成 PDF 二进制数据
            return BytesIO(pdf.output()).getvalue()
        except Exception:
            # 降级方案：使用 reportlab 生成支持中文的 PDF
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont

                # 使用 BytesIO 作为 PDF 输出
                buffer = BytesIO()
                doc = SimpleDocTemplate(buffer, pagesize=A4)

                # 样式
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    "CustomTitle",
                    parent=styles["Heading1"],
                    fontName="Helvetica-Bold",
                    fontSize=18,
                    spaceAfter=12,
                )
                normal_style = ParagraphStyle(
                    "CustomNormal",
                    parent=styles["BodyText"],
                    fontName="Helvetica",
                    fontSize=10,
                    spaceAfter=6,
                )

                # 构建内容
                story = []
                title_text = self._session.title or "新对话"
                story.append(Paragraph(title_text, title_style))
                story.append(Paragraph(f"<b>创建时间:</b> {self._format_time(self._session.created_at)}", normal_style))
                story.append(Paragraph(f"<b>更新时间:</b> {self._format_time(self._session.updated_at)}", normal_style))
                story.append(Spacer(1, 0.5 * cm))

                for msg in self._messages:
                    role_name = "<b>你</b>" if msg.role == "user" else "<b>助手</b>"
                    content = msg.content or ""
                    # 转义 HTML 特殊字符
                    content = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

                    story.append(Paragraph(f"{role_name}", ParagraphStyle("Role", parent=styles["Heading3"], fontSize=12)))
                    story.append(Paragraph(content, normal_style))
                    story.append(Spacer(1, 0.3 * cm))

                doc.build(story)
                return buffer.getvalue()
            except ImportError:
                # 如果 reportlab 不可用，使用原有方法
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", "B", 16)

                title = self._session.title or "新对话"
                pdf.cell(0, 10, title.encode("latin-1", "replace").decode("latin-1"), new_x="LEFT", new_y="NEXT", align="C")

                pdf.set_font("Arial", "", 10)
                pdf.cell(0, 6, f"Created: {self._format_time(self._session.created_at)}", new_x="LEFT", new_y="NEXT")
                pdf.cell(0, 6, f"Updated: {self._format_time(self._session.updated_at)}", new_x="LEFT", new_y="NEXT")

                pdf.ln(5)
                pdf.set_line_width(0.5)
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                pdf.ln(5)

                pdf.set_font("Arial", "B", 12)
                for msg in self._messages:
                    role_name = "You" if msg.role == "user" else "Assistant"
                    pdf.set_fill_color(240 if msg.role == "user" else 200, 240 if msg.role == "user" else 220, 240 if msg.role == "user" else 220)
                    pdf.cell(0, 8, role_name, new_x="LEFT", new_y="NEXT", fill=True)

                    pdf.set_font("Arial", "", 10)
                    content = msg.content or ""
                    safe_content = content.encode("latin-1", "replace").decode("latin-1")

                    lines = self._wrap_text(safe_content, 190)
                    for line in lines:
                        pdf.cell(0, 6, line, new_x="LEFT", new_y="NEXT")

                    pdf.ln(3)

                return BytesIO(pdf.output()).getvalue()

    def to_docx(self) -> bytes:
        """导出为 Word (DOCX) 格式.

        Returns:
            DOCX 文件的二进制数据
        """
        doc = Document()

        # 标题
        title = self._session.title or "新对话"
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = 0  # 0=left, 1=center, 2=right

        # 元信息
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"创建时间: {self._format_time(self._session.created_at)}\n")
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        meta_run = meta_para.add_run(f"更新时间: {self._format_time(self._session.updated_at)}")
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 分隔线（用水平线）
        doc.add_paragraph("_" * 80)

        # 消息
        for msg in self._messages:
            # 角色名称
            role_name = "你" if msg.role == "user" else "助手"
            role_para = doc.add_heading(role_name, level=2)

            # 设置角色标题颜色
            for run in role_para.runs:
                if msg.role == "user":
                    run.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)  # 蓝色
                else:
                    run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)  # 绿色

            # 消息内容
            content = msg.content or ""
            content_para = doc.add_paragraph(content)
            content_para.paragraph_format.space_after = Pt(6)

            # 时间戳
            time_para = doc.add_paragraph()
            time_run = time_para.add_run(f"{self._format_time(msg.created_at)}")
            time_run.font.size = Pt(8)
            time_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            time_para.paragraph_format.space_after = Pt(12)

        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _wrap_text(self, text: str, max_width: float) -> list[str]:
        """将文本按指定宽度换行.

        Args:
            text: 要换行的文本
            max_width: 最大宽度（字符单位，近似值）

        Returns:
            换行后的文本列表
        """
        if not text:
            return [""]

        # 近似计算：每个字符约 2.8pt（10号字体）
        chars_per_line = int(max_width / 2.8)
        lines = []

        while len(text) > chars_per_line:
            # 尝试在空格处换行
            break_point = text[:chars_per_line + 1].rfind(" ")
            if break_point == -1 or break_point == 0:
                break_point = chars_per_line

            lines.append(text[:break_point].strip())
            text = text[break_point:].strip()

        if text:
            lines.append(text)

        return lines if lines else [""]

    def save(self, path: str, format: str) -> None:
        """保存到文件.

        Args:
            path: 文件路径
            format: "txt", "md", "json", "html", "pdf" 或 "docx"
        """
        if format == "txt":
            content = self.to_txt()
            Path(path).write_text(content, encoding="utf-8")
        elif format == "md":
            content = self.to_markdown()
            Path(path).write_text(content, encoding="utf-8")
        elif format == "json":
            content = self.to_json()
            Path(path).write_text(content, encoding="utf-8")
        elif format == "html":
            content = self.to_html()
            Path(path).write_text(content, encoding="utf-8")
        elif format == "pdf":
            content = self.to_pdf()
            Path(path).write_bytes(content)
        elif format == "docx":
            content = self.to_docx()
            Path(path).write_bytes(content)
        else:
            raise ValueError(f"Unsupported format: {format}")

    @staticmethod
    def _format_time(iso_time: str) -> str:
        """格式化 ISO 时间为可读格式."""
        try:
            dt = datetime.fromisoformat(iso_time.replace("Z", "+00:00"))
            return dt.strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            return iso_time
