"""导出功能测试."""
from datetime import datetime, timezone

import pytest

from src.app.exporter import ChatExporter
from src.persistence.models import Session, Message


@pytest.fixture
def sample_session() -> Session:
    """示例会话 fixture."""
    return Session(
        id="s1",
        title="Test Session",
        created_at=datetime.now(timezone.utc).isoformat(),
        updated_at=datetime.now(timezone.utc).isoformat(),
    )


@pytest.fixture
def sample_messages() -> list[Message]:
    """示例消息列表 fixture."""
    now = datetime.now(timezone.utc).isoformat()
    return [
        Message(id="m1", session_id="s1", role="user", content="Hello", created_at=now),
        Message(id="m2", session_id="s1", role="assistant", content="Hi there!", created_at=now),
    ]


class TestChatExporter:
    """ChatExporter 测试."""

    def test_to_markdown(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：导出为 Markdown 格式."""
        exporter = ChatExporter(sample_session, sample_messages)
        md = exporter.to_markdown()

        assert "# Test Session" in md
        assert "**创建时间**" in md
        assert "## 你" in md
        assert "Hello" in md
        assert "## 助手" in md
        assert "Hi there!" in md

    def test_to_json(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：导出为 JSON 格式."""
        exporter = ChatExporter(sample_session, sample_messages)
        json_str = exporter.to_json()

        import json
        data = json.loads(json_str)

        assert data["session"]["id"] == "s1"
        assert data["session"]["title"] == "Test Session"
        assert len(data["messages"]) == 2
        assert data["messages"][0]["role"] == "user"
        assert data["messages"][0]["content"] == "Hello"
        assert data["messages"][1]["role"] == "assistant"
        assert data["messages"][1]["content"] == "Hi there!"

    def test_save_markdown(self, sample_session: Session, sample_messages: list[Message], temp_dir) -> None:
        """测试：保存 Markdown 文件."""
        exporter = ChatExporter(sample_session, sample_messages)
        path = str(temp_dir / "export.md")
        exporter.save(path, "md")

        import os
        assert os.path.exists(path)
        content = open(path, encoding="utf-8").read()
        assert "# Test Session" in content

    def test_save_json(self, sample_session: Session, sample_messages: list[Message], temp_dir) -> None:
        """测试：保存 JSON 文件."""
        exporter = ChatExporter(sample_session, sample_messages)
        path = str(temp_dir / "export.json")
        exporter.save(path, "json")

        import os
        import json
        assert os.path.exists(path)
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        assert data["session"]["id"] == "s1"

    def test_empty_session(self, sample_session: Session) -> None:
        """测试：空会话导出."""
        exporter = ChatExporter(sample_session, [])
        md = exporter.to_markdown()
        json_str = exporter.to_json()

        assert "# Test Session" in md
        assert "## 你" not in md

        import json
        data = json.loads(json_str)
        assert data["messages"] == []

    def test_to_pdf_returns_bytes(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：导出为 PDF 格式返回字节数据."""
        exporter = ChatExporter(sample_session, sample_messages)
        pdf_bytes = exporter.to_pdf()

        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 0
        # PDF 文件以 %PDF- 开头
        assert pdf_bytes[:4] == b"%PDF"

    def test_to_pdf_with_empty_session(self, sample_session: Session) -> None:
        """测试：空会话导出为 PDF."""
        exporter = ChatExporter(sample_session, [])
        pdf_bytes = exporter.to_pdf()

        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 0
        assert pdf_bytes[:4] == b"%PDF"

    def test_save_pdf(self, sample_session: Session, sample_messages: list[Message], temp_dir) -> None:
        """测试：保存 PDF 文件."""
        exporter = ChatExporter(sample_session, sample_messages)
        path = str(temp_dir / "export.pdf")
        exporter.save(path, "pdf")

        import os
        assert os.path.exists(path)
        with open(path, "rb") as f:
            content = f.read()
        assert content[:4] == b"%PDF"

    def test_save_unsupported_format_raises_error(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：不支持的格式抛出异常."""
        exporter = ChatExporter(sample_session, sample_messages)

        with pytest.raises(ValueError, match="Unsupported format"):
            exporter.save("export.txt", "txt")

    def test_wrap_text_short(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：短文本换行."""
        exporter = ChatExporter(sample_session, sample_messages)
        lines = exporter._wrap_text("Short text", 190)

        assert len(lines) == 1
        assert lines[0] == "Short text"

    def test_wrap_text_long(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：长文本换行."""
        exporter = ChatExporter(sample_session, sample_messages)
        long_text = "This is a very long text that should be wrapped into multiple lines for testing purposes."
        lines = exporter._wrap_text(long_text, 50)

        assert len(lines) > 1

    def test_wrap_text_empty(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：空文本换行."""
        exporter = ChatExporter(sample_session, sample_messages)
        lines = exporter._wrap_text("", 190)

        assert lines == [""]


class TestChatExporterHTML:
    """ChatExporter HTML 导出测试 (v1.0.7)."""

    def test_to_html_returns_string(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：导出为 HTML 格式返回字符串."""
        exporter = ChatExporter(sample_session, sample_messages)
        html = exporter.to_html()

        assert isinstance(html, str)
        assert len(html) > 0

    def test_to_html_contains_doctype(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：HTML 包含 DOCTYPE 声明."""
        exporter = ChatExporter(sample_session, sample_messages)
        html = exporter.to_html()

        assert "<!DOCTYPE html>" in html
        assert "<html" in html
        assert "</html>" in html

    def test_to_html_contains_title(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：HTML 包含会话标题."""
        exporter = ChatExporter(sample_session, sample_messages)
        html = exporter.to_html()

        assert "Test Session" in html
        assert "<title>" in html

    def test_to_html_contains_messages(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：HTML 包含消息内容."""
        exporter = ChatExporter(sample_session, sample_messages)
        html = exporter.to_html()

        assert "Hello" in html
        assert "Hi there!" in html
        assert "message user" in html
        assert "message assistant" in html

    def test_to_html_contains_styles(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：HTML 包含样式定义."""
        exporter = ChatExporter(sample_session, sample_messages)
        html = exporter.to_html()

        assert "<style>" in html
        assert "</style>" in html
        assert "body {" in html
        assert "font-family:" in html

    def test_to_html_with_chinese_content(self) -> None:
        """测试：HTML 导出支持中文内容."""
        session = Session(
            id="s1",
            title="测试会话",
            created_at=datetime.now(timezone.utc).isoformat(),
            updated_at=datetime.now(timezone.utc).isoformat(),
        )
        messages = [
            Message(id="m1", session_id="s1", role="user", content="你好，世界！", created_at=datetime.now(timezone.utc).isoformat()),
            Message(id="m2", session_id="s1", role="assistant", content="你好！有什么可以帮助你的吗？", created_at=datetime.now(timezone.utc).isoformat()),
        ]

        exporter = ChatExporter(session, messages)
        html = exporter.to_html()

        assert "测试会话" in html
        assert "你好，世界！" in html
        assert "你好！有什么可以帮助你的吗？" in html
        # 检查是否正确转义
        assert "&lt;" not in html or "Hello" in html  # 只有特殊字符才转义

    def test_save_html(self, sample_session: Session, sample_messages: list[Message], temp_dir) -> None:
        """测试：保存 HTML 文件."""
        exporter = ChatExporter(sample_session, sample_messages)
        path = str(temp_dir / "export.html")
        exporter.save(path, "html")

        import os
        assert os.path.exists(path)
        with open(path, encoding="utf-8") as f:
            content = f.read()
        assert "<!DOCTYPE html>" in content
        assert "Test Session" in content

    def test_to_html_with_empty_session(self, sample_session: Session) -> None:
        """测试：空会话导出为 HTML."""
        exporter = ChatExporter(sample_session, [])
        html = exporter.to_html()

        assert "<!DOCTYPE html>" in html
        assert "Test Session" in html
        # 没有消息
        assert "message user" not in html
        assert "message assistant" not in html


class TestChatExporterChinese:
    """ChatExporter 中文内容测试 (v1.0.7)."""

    def test_to_markdown_with_chinese(self) -> None:
        """测试：Markdown 导出支持中文."""
        session = Session(
            id="s1",
            title="中文测试",
            created_at=datetime.now(timezone.utc).isoformat(),
            updated_at=datetime.now(timezone.utc).isoformat(),
        )
        messages = [
            Message(id="m1", session_id="s1", role="user", content="你好，AI助手！", created_at=datetime.now(timezone.utc).isoformat()),
            Message(id="m2", session_id="s1", role="assistant", content="你好！请问有什么可以帮助你的？", created_at=datetime.now(timezone.utc).isoformat()),
        ]

        exporter = ChatExporter(session, messages)
        md = exporter.to_markdown()

        assert "# 中文测试" in md
        assert "你好，AI助手！" in md
        assert "你好！请问有什么可以帮助你的？" in md

    def test_to_json_with_chinese(self) -> None:
        """测试：JSON 导出支持中文（ensure_ascii=False）."""
        session = Session(
            id="s1",
            title="中文测试",
            created_at=datetime.now(timezone.utc).isoformat(),
            updated_at=datetime.now(timezone.utc).isoformat(),
        )
        messages = [
            Message(id="m1", session_id="s1", role="user", content="你好", created_at=datetime.now(timezone.utc).isoformat()),
        ]

        exporter = ChatExporter(session, messages)
        json_str = exporter.to_json()

        import json
        data = json.loads(json_str)

        assert data["session"]["title"] == "中文测试"
        assert data["messages"][0]["content"] == "你好"
        # JSON 字符串应直接包含中文，不是 Unicode 转义
        assert "你好" in json_str
        assert "\\u4f60\\u597d" not in json_str  # 不是 Unicode 转义


class TestChatExporterDOCX:
    """ChatExporter DOCX 导出测试 (v1.0.9)."""

    def test_to_docx_returns_bytes(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：导出为 DOCX 格式返回字节数据."""
        exporter = ChatExporter(sample_session, sample_messages)
        docx_bytes = exporter.to_docx()

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        # DOCX 文件（ZIP 格式）以 PK\x03\x04 开头
        assert docx_bytes[:4] == b"PK\x03\x04"

    def test_to_docx_contains_content(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：DOCX 包含正确的文本内容."""
        from docx import Document

        exporter = ChatExporter(sample_session, sample_messages)
        docx_bytes = exporter.to_docx()

        # 从字节数据加载文档
        from io import BytesIO
        doc = Document(BytesIO(docx_bytes))

        # 检查标题
        assert doc.paragraphs[0].text == "Test Session"

        # 检查是否包含消息内容
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Hello" in full_text
        assert "Hi there!" in full_text

    def test_to_docx_with_chinese_content(self) -> None:
        """测试：DOCX 导出支持中文内容."""
        from docx import Document
        from io import BytesIO

        session = Session(
            id="s1",
            title="中文测试",
            created_at=datetime.now(timezone.utc).isoformat(),
            updated_at=datetime.now(timezone.utc).isoformat(),
        )
        messages = [
            Message(id="m1", session_id="s1", role="user", content="你好，世界！", created_at=datetime.now(timezone.utc).isoformat()),
            Message(id="m2", session_id="s1", role="assistant", content="你好！有什么可以帮助你的吗？", created_at=datetime.now(timezone.utc).isoformat()),
        ]

        exporter = ChatExporter(session, messages)
        docx_bytes = exporter.to_docx()

        # 加载文档并检查中文内容
        doc = Document(BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "中文测试" in full_text
        assert "你好，世界！" in full_text
        assert "你好！有什么可以帮助你的吗？" in full_text

    def test_to_docx_with_empty_session(self, sample_session: Session) -> None:
        """测试：空会话导出为 DOCX."""
        exporter = ChatExporter(sample_session, [])
        docx_bytes = exporter.to_docx()

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
        assert docx_bytes[:4] == b"PK\x03\x04"

    def test_save_docx(self, sample_session: Session, sample_messages: list[Message], temp_dir) -> None:
        """测试：保存 DOCX 文件."""
        import os

        exporter = ChatExporter(sample_session, sample_messages)
        path = str(temp_dir / "export.docx")
        exporter.save(path, "docx")

        assert os.path.exists(path)
        with open(path, "rb") as f:
            content = f.read()
        assert content[:4] == b"PK\x03\x04"

    def test_to_docx_structure(self, sample_session: Session, sample_messages: list[Message]) -> None:
        """测试：DOCX 文档结构正确（标题、元信息、消息）。"""
        from docx import Document
        from io import BytesIO

        exporter = ChatExporter(sample_session, sample_messages)
        docx_bytes = exporter.to_docx()

        doc = Document(BytesIO(docx_bytes))

        # 检查文档结构
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # 应该包含标题
        assert "Test Session" in full_text

        # 应该包含时间戳信息
        assert "创建时间:" in full_text or "Created:" in full_text

        # 应该包含消息
        assert "Hello" in full_text
        assert "Hi there!" in full_text
